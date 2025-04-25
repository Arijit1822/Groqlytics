from django.shortcuts import render, HttpResponse
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt 
from django.http import JsonResponse
import psutil
import datetime
import subprocess  # To get active window
import time
import threading
from docx import Document
from docx.shared import Inches
import os
from groq import Groq
from django.core.cache import cache
from datetime import datetime
import psutil
import threading
from django.views.decorators.csrf import ensure_csrf_cookie
from django.http import JsonResponse
from django.core.cache import cache
from collections import deque

# Globals
app_history = []
start_time = None
current_app = None
network_history = {
    'download': deque(maxlen=10),
    'upload': deque(maxlen=10),
    'total_download': 0,
    'total_upload': 0,
    'last_check': psutil.net_io_counters()
}

# Your groq key
GROQ_API_KEY = "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"

# Initialize Groq client
groq_client = Groq(api_key=GROQ_API_KEY)







def get_active_window():
    try:
        # Windows
        import win32gui
        import win32process
        import psutil
        
        hwnd = win32gui.GetForegroundWindow()
        _, pid = win32process.GetWindowThreadProcessId(hwnd)
        if pid:
            app_name = psutil.Process(pid).name()
            return app_name
        return "N/A"
    except Exception as e:
        print(f"Error getting active window: {e}")
        return "N/A"

def suggest_improvement(app_name, time_spent):
    try:
        app_name = app_name.lower()
        prompt = f"App: {app_name}, Time: {time_spent} seconds. Suggest a productivity tip (max 50 words)."
        
        if "chrome" in app_name:
            prompt += " Consider if it's for work or distraction."
        elif "word" in app_name or "excel" in app_name:
            prompt += " Focus on completing tasks efficiently."
        
        chat_completion = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",  # Using a different model since mixtral is having issues
            temperature=0,
            max_tokens=100,
            top_p=1,
            stream=False,
        )
        return chat_completion.choices[0].message.content.strip() or "Stay productive!"
    except Exception as e:
        print(f"Suggestion error: {str(e)}")
        return "Stay focused on your tasks!"

def track_active_app():
    global current_app, start_time, app_history
    while True:
        active_app = get_active_window()
        current_time = time.time()
        
        if active_app != current_app:
            # Log the previous app's time before switching
            if current_app and start_time:
                elapsed_time = round(current_time - start_time)
                if elapsed_time > 0:
                    app_history.append((current_app, elapsed_time))
                    print(f"Added to history: {current_app} - {elapsed_time}s")
            
            # Reset timer for new app
            current_app = active_app
            start_time = current_time
        time.sleep(1)

app_tracker = threading.Thread(target=track_active_app, daemon=True)
app_tracker.start()

@login_required
@ensure_csrf_cookie
def dashboard(request):
    # Calculate current time spent
    current_time_spent = round(time.time() - start_time) if start_time else 0
    # Get existing context
    context = {
        'active_app': current_app if current_app else 'N/A',
        'time_spent': current_time_spent,
        'suggestion': suggest_improvement(current_app, current_time_spent) if current_app else "",
        'app_history': app_history
    }
    
    
    return render(request, 'dashboard.html', context)

@csrf_exempt
def reset_history(request):
    if request.method == 'POST':
        print("✅ reset_history view was called!") 
        global app_history, start_time, current_app
        try:
            print("Resetting history...") # Debug log
            app_history.clear()  # Clear the list instead of reassigning
            start_time = None
            current_app = None
            
            return JsonResponse({
                'status': 'reset',
                'message': 'History reset successfully'
            })
        except Exception as e:
            print(f"Reset error: {str(e)}") # Debug log
            return JsonResponse({
                'status': 'error',
                'message': str(e)
            }, status=500)
    return JsonResponse({
        'status': 'error',
        'message': 'Invalid request method'
    }, status=405)

def save_logs_pdf(request):
    global app_history
    document = Document()
    document.add_heading('App Usage History', 0)
    for app, duration in app_history:
        document.add_paragraph(f'{app}: {duration} seconds')
    filename = "app_usage_history.docx"
    document.save(filename)

    # Convert docx to pdf using soffice
    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', filename], check=True)
        pdf_filename = filename.replace('.docx', '.pdf')

        with open(pdf_filename, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
        os.remove(filename)
        os.remove(pdf_filename)
        return response
    except Exception:
        return HttpResponse(f"Error converting to PDF: Add soffice to PATH to use this feature")
    
#Ignore this......this is not doing anything
def get_network_usage():
    global network_history
    current = psutil.net_io_counters()
    last = network_history['last_check']
    
    # Calculate speeds
    download_speed = (current.bytes_recv - last.bytes_recv) / 1024 / 1024  # MB/s
    upload_speed = (current.bytes_sent - last.bytes_sent) / 1024 / 1024    # MB/s
    
    # Update totals
    network_history['total_download'] += download_speed
    network_history['total_upload'] += upload_speed
    
    # Update history
    network_history['download'].append(round(download_speed, 2))
    network_history['upload'].append(round(upload_speed, 2))
    
    # Update last check
    network_history['last_check'] = current
    
    return {
        'download_speed': round(download_speed, 2),
        'upload_speed': round(upload_speed, 2),
        'total_download': round(network_history['total_download'], 2),
        'total_upload': round(network_history['total_upload'], 2),
        'history_download': list(network_history['download']),
        'history_upload': list(network_history['upload'])
    }



@login_required
def system_api(request):
    cpu_freq = int(psutil.cpu_freq().current) if psutil.cpu_freq() else 0
    cpu_cores = psutil.cpu_count(logical=False)
    cpu_usage = psutil.cpu_percent(interval=0.1)
    ram = psutil.virtual_memory()
    disk = psutil.disk_usage('/')
    now = datetime.now()
    network_data = get_network_usage()
    
    # Format history data
    formatted_history = [
        {"app": app, "time_spent": duration}
        for app, duration in app_history
    ]
    
    current_time_spent = round(time.time() - start_time) if start_time else 0
    
    return JsonResponse({
        'cpu_freq': cpu_freq,
        'cpu_cores': cpu_cores,
        'cpu_usage': cpu_usage,
        'total_storage': round(disk.total/(1024**3),2),
        'total_disk': round(disk.total/(1024**3),2),
        'ram_used': round(ram.used/(1024**3),2),
        'ram_total': round(ram.total/(1024**3),2),
        'disk_used': round(disk.used/(1024**3),2),
        'disk_total': round(disk.total/(1024**3),2),
        'time': now.strftime('%Y-%m-%d %H:%M:%S'),
        'active_app': current_app if current_app else 'N/A',
        'time_spent': current_time_spent,
        'suggestion': suggest_improvement(current_app, current_time_spent) if current_app else 'No suggestion',
        'history': formatted_history, # Updated format
        'network': network_data
    })


